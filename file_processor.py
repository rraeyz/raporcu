import os
import io
import docx
import fitz  # PyMuPDF
import tempfile

class FileProcessor:
    """Dosya işleme sınıfı"""
    def __init__(self):
        pass
    
    def extract_text_from_file(self, file_path):
        """Farklı dosya formatlarından metin çıkarır"""
        try:
            # Dosya uzantısını al
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # TXT dosyası
            if file_ext == '.txt':
                return self._extract_from_txt(file_path)
            
            # PDF dosyası
            elif file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            
            # DOCX dosyası
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            
            # Desteklenmeyen format
            else:
                print(f"Desteklenmeyen dosya formatı: {file_ext}")
                return None
                
        except Exception as e:
            print(f"Dosyadan metin çıkarılırken hata: {e}")
            return None
    
    def _extract_from_txt(self, file_path):
        """TXT dosyasından metin çıkarır"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # UTF-8 ile açılamazsa farklı kodlamalar dene
            for encoding in ['latin-1', 'cp1254', 'iso-8859-9']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # Hiçbir kodlama ile açılamazsa
            print("Dosya kodlaması tespit edilemedi.")
            return None
        except Exception as e:
            print(f"TXT dosyası okunurken hata: {e}")
            return None
    
    def _extract_from_pdf(self, file_path):
        """PDF dosyasından metin çıkarır"""
        try:
            text = ""
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
            return text
        except Exception as e:
            print(f"PDF dosyası okunurken hata: {e}")
            return None
    
    def _extract_from_docx(self, file_path):
        """DOCX dosyasından metin çıkarır"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Paragrafları birleştir
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"DOCX dosyası okunurken hata: {e}")
            return None
    
    def _parse_markdown(self, text_content):
        """Markdown metnini parçalara ayırır"""
        # Başlık ve alt başlıkları tespit et
        import re
        content_parts = []
        current_lines = []
        in_code_block = False
        
        for line in text_content.split('\n'):
            # Kod bloğu başlangıç/bitiş kontrolü
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                current_lines.append(line)
                continue
                
            # Kod bloğu içindeyse direk ekle
            if in_code_block:
                current_lines.append(line)
                continue
                
            # Markdown başlık kontrolü
            if line.strip().startswith('#') and not in_code_block:
                if current_lines:
                    content_parts.append('\n'.join(current_lines))
                    current_lines = []
                
                # Başlık seviyesini belirle (h1, h2, h3)
                heading_level = len(re.match(r'^#+', line).group())
                title_text = line.strip('#').strip()
                
                content_parts.append({
                    'type': f'h{heading_level}',
                    'content': title_text
                })
            else:
                current_lines.append(line)
        
        if current_lines:
            content_parts.append('\n'.join(current_lines))
            
        return content_parts
            
    def save_as_docx(self, text_content, output_path, metadata=None, config_instance=None):
        """Metni DOCX dosyası olarak kaydeder ve şablona göre formatlar"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import re
        import os
        
        try:
            # Metadata kontrolü
            if metadata is None:
                metadata = {}
            
            # Markdown'ı parçala
            content_parts = self._parse_markdown(text_content)
            
            # Yeni Word belgesi oluştur
            doc = Document()
            
            # Ana uygulamadan gelen config örneğini kullan veya yoksa yeni oluştur
            if config_instance:
                config = config_instance
                print("FileProcessor: Ana uygulamadan gelen config örneği kullanılıyor")
            else:
                # Dikkat: Yeni bir Config örneği oluşturmak, şablonların düzgün yüklenmemesine neden olabilir
                from config import Config
                print("FileProcessor: Yeni bir Config örneği oluşturuluyor (senkron modda)")
                # Senkron modda yükle (async_load=False)
                config = Config(async_load=False)
            
            template = config.get_active_template()
            
            # Debug: Şablon bilgilerini yazdır
            print(f"Aktif şablon ID: {config.active_template}")
            print(f"Şablon adı: {template.get('name', 'Bilinmiyor')}")
            print(f"Şablon içeriği: {template.keys()}")
            
            # Belge stillerini ayarla
            styles = doc.styles
            
            # Şablondan stil bilgilerini al
            style_settings = template.get('settings', {})
            main_title_settings = style_settings.get('main_title', {})
            headings_settings = style_settings.get('headings', {})
            subheadings_settings = style_settings.get('subheadings', {})
            text_settings = style_settings.get('text', {})
            table_settings = style_settings.get('tables', {})
            image_settings = style_settings.get('images', {})
            
            # Belge fontunu ve stillerini düzenle
            style_normal = styles['Normal']
            style_normal.font.name = text_settings.get('font_name', 'Times New Roman')
            style_normal.font.size = Pt(text_settings.get('font_size', 12))
            style_normal.paragraph_format.space_after = Pt(8)
            
            # Başlık stillerini düzenle
            for i in range(1, 5):
                style_name = f'Heading {i}'
                if style_name in styles:
                    style = styles[style_name]
                    
                    # Ana başlık ayarları (Heading 1 için)
                    if i == 1:
                        style.font.name = main_title_settings.get('font_name', 'Times New Roman')
                        style.font.bold = main_title_settings.get('bold', True)
                        style.font.size = Pt(main_title_settings.get('font_size', 12))
                        style.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
                        alignment = main_title_settings.get('alignment', 'center')
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                    # Alt başlık ayarları (Heading 2-4 için)
                    else:
                        settings = headings_settings if i == 2 else subheadings_settings
                        style.font.name = settings.get('font_name', 'Times New Roman')
                        style.font.bold = settings.get('bold', True)
                        style.font.size = Pt(settings.get('font_size', 12))
                        style.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
                        alignment = settings.get('alignment', 'left')
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                    
                    style.paragraph_format.space_before = Pt(12)
                    style.paragraph_format.space_after = Pt(8)
                    
                    # Başlık altındaki çizgiyi kaldır
                    if hasattr(style.paragraph_format, 'border_bottom'):
                        style.paragraph_format.border_bottom.width = 0
            
            # Özel kapak sayfası stili
            if 'Title' in styles:
                title_style = styles['Title']
                title_style.font.name = main_title_settings.get('font_name', 'Times New Roman')
                title_style.font.bold = main_title_settings.get('bold', True)
                title_style.font.size = Pt(main_title_settings.get('font_size', 12))
                title_style.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
                alignment = main_title_settings.get('alignment', 'center')
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                title_style.paragraph_format.space_after = Pt(12)
            
            # Tablo başlık stili
            table_header_style = styles.add_style('Table Header', WD_STYLE_TYPE.PARAGRAPH)
            table_header_style.font.bold = True
            table_header_style.font.size = Pt(text_settings.get('font_size', 12))
            table_header_style.font.name = text_settings.get('font_name', 'Times New Roman')
            table_header_style.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
            
            # Görsel başlık stili
            caption_style = styles.add_style('Image Caption', WD_STYLE_TYPE.PARAGRAPH)
            caption_style.font.italic = True
            caption_style.font.size = Pt(text_settings.get('font_size', 12))
            caption_style.font.name = text_settings.get('font_name', 'Times New Roman')
            caption_style.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
            caption_alignment = image_settings.get('caption_alignment', 'left')
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if caption_alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT
            
            # İçeriği Word'e ekle
            for part in content_parts:
                if isinstance(part, dict):  # Başlık
                    heading_type = part['type']
                    heading_text = part['content']
                    
                    level = int(heading_type[1])  # h1, h2, h3, ... seviyesi
                    
                    # Başlığı ekle
                    heading = doc.add_heading(level=level)
                    heading_run = heading.add_run(heading_text)
                    
                    # Başlık stilini ayarla
                    if level == 1:
                        settings = main_title_settings
                    elif level == 2:
                        settings = headings_settings
                    else:
                        settings = subheadings_settings
                    
                    # Başlık yazı tipi ayarları
                    heading_run.font.name = settings.get('font_name', 'Times New Roman')
                    heading_run.font.bold = settings.get('bold', True)
                    heading_run.font.size = Pt(settings.get('font_size', 12))
                    
                    # Başlık hizalama
                    alignment = settings.get('alignment', 'left' if level > 1 else 'center')
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                    
                else:  # Normal metin
                    # Metnin boş satırlarını parçala
                    paragraphs = part.split('\n')
                    
                    for para_text in paragraphs:
                        if para_text.strip():  # Boş olmayan satırlar
                            # Paragraf ekle
                            paragraph = doc.add_paragraph()
                            
                            # Markdown formatlamaları
                            # Kalın metin
                            bold_parts = re.split(r'(\*\*.*?\*\*)', para_text)
                            
                            for bold_part in bold_parts:
                                if bold_part.startswith('**') and bold_part.endswith('**'):
                                    # Kalın metin
                                    run = paragraph.add_run(bold_part.strip('**'))
                                    run.bold = True
                                    run.font.name = text_settings.get('font_name', 'Times New Roman')
                                    run.font.size = Pt(text_settings.get('font_size', 12))
                                else:
                                    # Normal metin
                                    if bold_part.strip():
                                        run = paragraph.add_run(bold_part)
                                        run.font.name = text_settings.get('font_name', 'Times New Roman')
                                        run.font.size = Pt(text_settings.get('font_size', 12))
                            
                            # Paragraf hizalama
                            alignment = text_settings.get('alignment', 'left')
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if alignment == 'justify' else (
                                WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT)
            
            # Dosyayı kaydet
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"DOCX kaydedilirken hata: {e}")
            import traceback
            traceback.print_exc()
            return False

    def convert_docx_to_pdf(self, docx_path, pdf_path):
        """DOCX dosyasını PDF'e dönüştürür"""
        try:
            # PDF oluşturmak için alternatif bir yöntem kullanıyoruz
            print(f"DOCX dosyası PDF'e dönüştürülüyor: {docx_path} -> {pdf_path}")
            
            # Windows'ta LibreOffice/MSOffice COM API kullanarak dönüşüm
            import sys
            import subprocess
            
            # Microsoft Word kullanarak dönüşüm (Windows için)
            if sys.platform == 'win32':
                try:
                    import win32com.client
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    # Tam dosya yolları
                    doc_path = os.path.abspath(docx_path)
                    pdf_path_abs = os.path.abspath(pdf_path)
                    
                    # DOCX'i aç
                    doc = word.Documents.Open(doc_path)
                    # PDF olarak kaydet (17 = PDF formatı)
                    doc.SaveAs(pdf_path_abs, FileFormat=17)
                    doc.Close()
                    word.Quit()
                    return True
                except Exception as word_error:
                    print(f"MS Word dönüşüm hatası: {word_error}")
                    # MS Word ile olmadıysa alternatif yöntem dene
            
            # Alternatif yöntem: Komut satırı araçlarını kullan
            try:
                # LibreOffice kullanarak dönüştür
                if sys.platform in ['linux', 'darwin']:
                    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", 
                                    "--outdir", os.path.dirname(pdf_path), docx_path], check=True)
                    return True
                else:
                    # PDF dönüşümü yapılamadı uyarısı
                    print("PDF dönüşümü için uygun bir araç bulunamadı.")
                    print("LibreOffice, Microsoft Word veya docx2pdf kurulu değil.")
                    return False
            except Exception as cmd_error:
                print(f"Komut satırı dönüşüm hatası: {cmd_error}")
                return False
            
        except Exception as e:
            print(f"DOCX to PDF dönüşüm hatası: {e}")
            return False

    def save_as_pdf(self, text_content, output_path, metadata=None, config_instance=None):
        """Metni PDF dosyası olarak kaydeder"""
        try:
            # Önce DOCX olarak kaydet
            temp_dir = tempfile.gettempdir()
            temp_docx = os.path.join(temp_dir, "temp_report.docx")
            
            # DOCX dosyasını oluştur - aynı config örneğini ilet
            self.save_as_docx(text_content, temp_docx, metadata, config_instance)
            
            # DOCX'i PDF'e dönüştür
            self.convert_docx_to_pdf(temp_docx, output_path)
            
            # Geçici dosyayı sil
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
                
            return True
        except Exception as e:
            print(f"PDF kaydedilirken hata: {e}")
            import traceback
            traceback.print_exc()
            return False
